# backend/documents/generator.py
"""
Генератор PDF и Excel документов
"""

import os
import subprocess
from typing import Optional
from pathlib import Path
from jinja2 import Template, Environment, FileSystemLoader
from datetime import date

from .models import Document
from backend.utils.latex import find_xelatex as _find_xelatex


class DocumentGenerator:
    """Генератор документов (PDF через LaTeX, Excel через openpyxl)"""

    def __init__(self, templates_dir: str = "backend/documents/templates"):
        """
        Args:
            templates_dir: Директория с шаблонами
        """
        self.templates_dir = Path(templates_dir)
        self.latex_dir = self.templates_dir / "latex"
        self.excel_dir = self.templates_dir / "excel"
        self.output_dir = Path("backend/static/generated")
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Создаём под директории
        (self.output_dir / "pdf").mkdir(exist_ok=True)
        (self.output_dir / "excel").mkdir(exist_ok=True)
        (self.output_dir / "temp").mkdir(exist_ok=True)

        # Jinja2 environment для LaTeX
        self.latex_env = Environment(
            loader=FileSystemLoader(str(self.latex_dir)),
            block_start_string='\\BLOCK{',
            block_end_string='}',
            variable_start_string='\\VAR{',
            variable_end_string='}',
            comment_start_string='\\#{',
            comment_end_string='}',
            line_statement_prefix='%%',
            line_comment_prefix='%#',
            trim_blocks=True,
            autoescape=False,
        )

    # ==========================================
    # ГЕНЕРАЦИЯ PDF
    # ==========================================

    def generate_pdf(self, document: Document) -> str:
        """
        Сгенерировать PDF для документа

        Args:
            document: Объект Document

        Returns:
            str: Путь к сгенерированному PDF (относительно static/)
        """
        # Получаем имя шаблона (fallback для work_plan)
        template_name = document.doc_type.latex_template_name
        if not template_name:
            code = getattr(document.doc_type, "code", "")
            if code == "work_plan":
                template_name = "work_plan.tex"
            else:
                raise ValueError(f"Для типа '{document.doc_type.code}' не указан LaTeX шаблон")

        # Готовим данные для шаблона
        context = self._prepare_context(document)

        # Рендерим LaTeX
        latex_source = self._render_latex_template(template_name, context)

        # Сохраняем исходник в документ
        document.latex_source = latex_source

        # Компилируем в PDF
        pdf_path = self._compile_latex_to_pdf(document, latex_source)

        # Сохраняем путь в документ
        document.pdf_filename = f"generated/pdf/{Path(pdf_path).name}"

        return document.pdf_filename

    def _prepare_context(self, document: Document) -> dict:
        """
        Подготовить контекст для шаблона

        Args:
            document: Объект Document

        Returns:
            dict: Контекст для Jinja2
        """
        # Базовый контекст
        context = {
            "doc": document,
            "doc_number": document.doc_number,
            "doc_type": document.doc_type,
            "well": document.well,
            "items": document.items,
            "period_start": document.period_start,
            "period_end": document.period_end,
            "metadata": document.metadata,
        }

        # Добавляем метаданные
        if document.metadata:
            context.update(document.metadata)

        # Специфичная обработка для акта расхода реагентов
        if document.doc_type.code == 'reagent_expense':
            context.update(self._prepare_reagent_expense_context(document))

        # Финансовый акт и счёт-фактура (одинаковый контекст, разные шаблоны)
        if document.doc_type.code in ('financial_act', 'financial_invoice'):
            context.update(self._prepare_financial_act_context(document))

        # План работ (work_plan) — Таблица 1.1 + подписанты
        if document.doc_type.code == 'work_plan':
            context.update(self._prepare_work_plan_context(document))

        return context

    # ==========================================
    # ФИНАНСОВЫЙ АКТ (общий контекст для .docx и PDF — инвариант PARITY)
    # ==========================================

    @staticmethod
    def _fin_fmt(x) -> str:
        """Число → '1 234 567,89' (пробел-разделитель тысяч, запятая-десятичная)."""
        from decimal import Decimal
        d = Decimal(str(x or 0)).quantize(Decimal("0.01"))
        s = f"{d:,.2f}"  # '1,234,567.89'
        return s.replace(",", " ").replace(".", ",")

    def _prepare_financial_act_context(self, document: Document) -> dict:
        """Единый источник данных финансового акта (для .docx и LaTeX)."""
        m = document.meta or {}
        order = ["adaptation", "optimization", "foam_dosing", "inhibitor_dosing"]
        titles = {
            "adaptation": "Адаптация / Adaptation",
            "optimization": "Оптимизация (дни/дней в месяце × ежемесячный платёж) / Optimization",
            "foam_dosing": "Дозирование пенных реагентов / Foam reagent dosing",
            "inhibitor_dosing": "Дозирование ингибирующих реагентов / Inhibitor reagent dosing",
        }
        buckets: dict[str, list] = {k: [] for k in order}
        for it in document.items:
            buckets.setdefault(it.work_group or "", []).append(it)

        groups = []
        for key in order:
            rows = buckets.get(key) or []
            if not rows:
                continue
            out = []
            for n, it in enumerate(rows, start=1):
                out.append({
                    "n": n,
                    "well": it.well_number or "",
                    "period": it.period_label or "",
                    "unit": it.unit or "",
                    "qty": it.quantity or 0,
                    "price": self._fin_fmt(it.price_per_unit),
                    "amount": self._fin_fmt(it.amount),
                    "vat": self._fin_fmt(it.vat_amount),
                    "total": self._fin_fmt(it.amount_with_vat),
                })
            groups.append({"title": titles.get(key, key), "rows": out})

        # Плоский список строк для docxtpl-цикла: название — только в первой строке группы
        flat = []
        for g in groups:
            for i, row in enumerate(g["rows"]):
                flat.append({**row, "name": g["title"] if i == 0 else ""})

        return {
            "rows": flat,
            "act_no": m.get("act_no", 1),
            "invoice_no": m.get("invoice_no", ""),
            "act_date": (document.period_end.strftime("%d.%m.%Y") if document.period_end else ""),
            "contract_ref": m.get("contract_ref", ""),
            "period_from": document.period_start.strftime("%d.%m.%Y") if document.period_start else "",
            "period_to": document.period_end.strftime("%d.%m.%Y") if document.period_end else "",
            "fin_groups": groups,
            "total_amount": self._fin_fmt(m.get("total_amount")),
            "total_vat": self._fin_fmt(m.get("total_vat")),
            "total_with_vat": self._fin_fmt(m.get("total_with_vat")),
            "words_ru": m.get("total_with_vat_words_ru", ""),
            "words_en": m.get("total_with_vat_words_en", ""),
            "vat_words_ru": m.get("total_vat_words_ru", ""),
            "vat_words_en": m.get("total_vat_words_en", ""),
            "wells_str": ", ".join(m.get("wells", [])),
            **self._decision_context(m),
            **self._signatory_context(m),
        }

    @staticmethod
    def _decision_context(m: dict) -> dict:
        """Строки решений: продолжить (по умолч. все) / прекратить (выбранные)."""
        cont = m.get("continue_wells", m.get("wells", []))
        stop = m.get("stop_wells", [])

        cc = m.get("continue_clause", "3.9")
        sc = m.get("stop_clause", "3.17")

        def phrase(wells, verb, clause):
            if not wells:
                return ""
            head = "На скважине №" if len(wells) == 1 else "На скважинах №№ "
            return f"{head}{', '.join(wells)} {verb} в соответствии с пунктом {clause} Договора."

        def phrase_en(wells, verb, clause):
            if not wells:
                return ""
            head = "At well #" if len(wells) == 1 else "At wells #"
            return f"{head}{', '.join(wells)} {verb} in accordance with clause {clause} of the Contract."

        return {
            "decision_continue": phrase(cont, "продолжить работы по оптимизации/обслуживанию", cc),
            "decision_stop": phrase(stop, "прекратить работы", sc),
            "decision_continue_en": phrase_en(cont, "continue optimization/maintenance work", cc),
            "decision_stop_en": phrase_en(stop, "cease works", sc),
        }

    @staticmethod
    def _signatory_context(m: dict) -> dict:
        """Из header_sigs/sign_sigs → строки шапки (по сторонам) и списки подписей."""
        header = m.get("header_sigs", [])
        sign = m.get("sign_sigs", [])

        def hstr(side, lang):
            key_p, key_n = (f"position_{lang}", f"name_{lang}")
            parts = [f"{s.get(key_p,'')} {s.get(key_n,'')}".strip()
                     for s in header if s.get("side") == side]
            return ", ".join(p for p in parts if p)

        return {
            "header_contractor": hstr("contractor", "ru"),
            "header_customer": hstr("customer", "ru"),
            "header_contractor_en": hstr("contractor", "en"),
            "header_customer_en": hstr("customer", "en"),
            "sign_contractor": [s for s in sign if s.get("side") == "contractor"],
            "sign_customer": [s for s in sign if s.get("side") == "customer"],
        }

    # ==========================================
    # ПЛАН РАБОТ (docx-шаблон + картинки печать/подпись)
    # ==========================================

    # Ассеты по умолчанию (можно переопределить путём в meta: seal_path / signature_path)
    DEFAULT_SEAL_PATH = "backend/static/img/unitool_seal.png"
    DEFAULT_SIGNATURE_PATH = "backend/static/img/verba_signature.png"

    def _prepare_work_plan_context(self, document: Document) -> dict:
        """Текстовый контекст «Плана работ»: Таблица 1.1 + подписанты.
        Печать/подпись накладываются ЗА текст отдельно (`_overlay_seal_signature`)."""
        from backend.services.work_plan_service import TABLE11_KEYS, DEFAULT_SIGS

        m = document.meta or {}
        t = {k: "" for k in TABLE11_KEYS}
        t.update(m.get("table11", {}) or {})

        ctx = {"t": t, "well_no": t.get("well_no", "")}
        for k, default in DEFAULT_SIGS.items():
            ctx[k] = m.get(k) or default
        return ctx

    @staticmethod
    def _overlay_seal_signature(path: str, document: Document) -> None:
        """Наложить печать (4 см) и подпись ЗА текст блока УТВЕРЖДАЮ (правая ячейка
        шапки). Плавающие якоря behindDoc, привязка к абзацу «ФИО» — двигаются
        вместе с блоком. Пропорции сохраняются."""
        from docx import Document as Docx
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from PIL import Image as PILImage

        m = document.meta or {}
        EMU_MM = 36000  # 1 мм
        seal_p = (m.get("seal_path") or DocumentGenerator.DEFAULT_SEAL_PATH)
        sig_p = (m.get("signature_path") or DocumentGenerator.DEFAULT_SIGNATURE_PATH)

        d = Docx(path)
        # правая ячейка шапки (таблица 0, строка 0, ячейка 1) → абзац «ФИО»
        cell = d.tables[0].rows[0].cells[1]
        # абзац с «_____ {ФИО}» — предпоследний непустой (перед строкой даты)
        name_par = None
        for p in cell.paragraphs:
            if "___" in p.text and "  " in p.text:
                name_par = p
        anchor_par = name_par or cell.paragraphs[-1]

        def add(img_path, w_mm, h_mm, x_mm, y_mm, zid, name):
            if not img_path or not os.path.exists(img_path):
                return
            rId, _ = d.part.get_or_add_image(img_path)
            cx, cy = int(w_mm * EMU_MM), int(h_mm * EMU_MM)
            x, y = int(x_mm * EMU_MM), int(y_mm * EMU_MM)
            xml = (
                f'<w:drawing {nsdecls("w", "wp", "a", "pic", "r")}>'
                f'<wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0" '
                f'simplePos="0" locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="{zid}">'
                f'<wp:simplePos x="0" y="0"/>'
                f'<wp:positionH relativeFrom="column"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
                f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{y}</wp:posOffset></wp:positionV>'
                f'<wp:extent cx="{cx}" cy="{cy}"/>'
                f'<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
                f'<wp:docPr id="{zid}" name="{name}"/><wp:cNvGraphicFramePr/>'
                f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
                f'<pic:pic><pic:nvPicPr><pic:cNvPr id="{zid}" name="{name}"/><pic:cNvPicPr/></pic:nvPicPr>'
                f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
                f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
                f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
                f'</a:graphicData></a:graphic></wp:anchor></w:drawing>'
            )
            anchor_par.add_run()._r.append(parse_xml(xml))

        # печать 40×40 мм; подпись — 28 мм по ширине, высота по пропорции
        if m.get("include_signature", True):
            try:
                w, h = PILImage.open(sig_p).size
            except Exception:
                w, h = 244, 137
            sig_w = 28.0
            sig_h = sig_w * h / w
            # подпись на подчёркивании: x=5мм, y=-5мм
            add(sig_p, sig_w, sig_h, x_mm=5, y_mm=-5, zid=101, name="signature")
        if m.get("include_seal", True):
            # печать накрывает подпись, не залезает на заголовок: y=-25мм
            add(seal_p, 40.0, 40.0, x_mm=10, y_mm=-25, zid=100, name="seal")

        d.save(path)

    def generate_docx(self, document: Document) -> str:
        """.docx по docxtpl-шаблону (реальный файл клиента). Иначе — программно (fallback)."""
        # шаблон по имени из типа; страховка — по коду (чтобы неверное имя в БД
        # не роняло в программный fallback)
        code = getattr(document.doc_type, "code", "")
        default_tpl = ("financial_invoice_template.docx" if code == "financial_invoice"
                       else "financial_act_template.docx")
        tpl_name = getattr(document.doc_type, "docx_template_name", None) or default_tpl
        tpl_path = self.templates_dir / "docx" / tpl_name
        if not tpl_path.exists():
            tpl_path = self.templates_dir / "docx" / default_tpl
        if tpl_path.exists():
            from docxtpl import DocxTemplate
            tpl = DocxTemplate(str(tpl_path))
            if code == "work_plan":
                ctx = self._prepare_work_plan_context(document)
            else:
                ctx = self._prepare_financial_act_context(document)
            tpl.render(ctx)
            base_name = document.doc_number.replace("/", "-")
            (self.output_dir / "docx").mkdir(exist_ok=True)
            out = self.output_dir / "docx" / f"{base_name}.docx"
            tpl.save(str(out))
            if code in ("financial_act", "financial_invoice"):
                self._style_work_table(str(out))  # жирные разделители + merge названий групп
            elif code == "work_plan":
                self._overlay_seal_signature(str(out), document)  # печать+подпись ЗА текст
            return f"generated/docx/{out.name}"
        return self._generate_docx_programmatic(document)

    @staticmethod
    def _style_work_table(path: str) -> None:
        """Постобработка таблицы работ: жирная линия перед каждой группой +
        вертикальное объединение ячейки «Наименование» по группам."""
        from docx import Document as Docx
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        def thick_top(tc):
            tcPr = tc.find(W + "tcPr")
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
            borders = tcPr.find(W + "tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
            for old in borders.findall(W + "top"):
                borders.remove(old)
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "18")
            top.set(qn("w:space"), "0"); top.set(qn("w:color"), "000000")
            borders.append(top)

        def set_vmerge(tc, restart):
            tcPr = tc.find(W + "tcPr")
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
            for old in tcPr.findall(W + "vMerge"):
                tcPr.remove(old)
            vm = OxmlElement("w:vMerge")
            if restart:
                vm.set(qn("w:val"), "restart")
            tcPr.append(vm)

        def close_bottom(tc):  # явная нижняя граница ячейки — закрывает merged-колонку,
            tcPr = tc.find(W + "tcPr")   # чтобы её боковые линии не «протекали» вниз
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
            borders = tcPr.find(W + "tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
            for old in borders.findall(W + "bottom"):
                borders.remove(old)
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
            borders.append(b)

        d = Docx(path)
        t = d.tables[1]
        rows = t.rows
        n = len(rows)
        # группы: строка-начало = непустая ячейка «Наименование» (физ. индекс 1); строки 2..n-2 — данные
        starts = []
        for i in range(2, n - 1):
            tcs = rows[i]._tr.findall(W + "tc")
            if len(tcs) > 1 and "".join(tcs[1].itertext()).strip():
                starts.append(i)
        for k, s in enumerate(starts):
            e = (starts[k + 1] - 1) if k + 1 < len(starts) else (n - 2)
            for tc in rows[s]._tr.findall(W + "tc"):  # жирная верхняя линия у всей строки
                thick_top(tc)
            for idx in range(s, e + 1):               # vMerge названия по группе
                tcs = rows[idx]._tr.findall(W + "tc")
                if len(tcs) > 1:
                    set_vmerge(tcs[1], restart=(idx == s))
                    if idx == s:                       # центрировать название по вертикали
                        tcPr = tcs[1].find(W + "tcPr")
                        for old in tcPr.findall(W + "vAlign"):
                            tcPr.remove(old)
                        va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "center")
                        tcPr.append(va)
                    if idx == e:                       # закрыть нижнюю границу merged-ячейки
                        close_bottom(tcs[1])           # (иначе её боковые линии текут вниз)

        if n >= 2:  # жирная линия перед «Всего» + закрыть низ таблицы явной границей
            for tc in rows[n - 1]._tr.findall(W + "tc"):
                thick_top(tc)
                close_bottom(tc)
        d.save(path)

    def _generate_docx_programmatic(self, document: Document) -> str:
        """Fallback: собрать .docx программно (без шаблона клиента)."""
        from docx import Document as Docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        ctx = self._prepare_financial_act_context(document)
        d = Docx()

        h = d.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run(f"Акт приёма-передачи выполненных работ № {ctx['act_no']} от {ctx['act_date']}\n"
                      f"Act of acceptance of rendered services # {ctx['act_no']} dated {ctx['act_date']}")
        r.bold = True

        d.add_paragraph(
            f"согласно Контракту № {ctx['contract_ref']} / in accordance with Contract no {ctx['contract_ref']}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        d.add_paragraph(
            "Мы, нижеподписавшиеся, представители заказчика СП ООО «Uz-Kor Gas Chemical» "
            "и представитель подрядчика ООО «UNITOOL», составили настоящий Акт о том, что "
            f"специалистами Исполнителя в период с {ctx['period_from']} по {ctx['period_to']} "
            "были выполнены следующие работы:"
        )

        headers = ["№", "Наименование работ / Name of work", "№ скв", "Период / Period",
                   "Ед. / unit", "К-во", "Цена за ед / Price", "Стоимость / Cost",
                   "НДС 12% / VAT", "С НДС / with VAT"]
        table = d.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for c, txt in zip(table.rows[0].cells, headers):
            c.paragraphs[0].add_run(txt).bold = True

        for g in ctx["fin_groups"]:
            gr = table.add_row().cells
            gr[0].merge(gr[len(headers) - 1])
            gr[0].paragraphs[0].add_run(g["title"]).bold = True
            for row in g["rows"]:
                cells = table.add_row().cells
                vals = [str(row["n"]), "", row["well"], row["period"], row["unit"],
                        str(row["qty"]), row["price"], row["amount"], row["vat"], row["total"]]
                for cell, v in zip(cells, vals):
                    cell.paragraphs[0].add_run(v)

        tr = table.add_row().cells
        tr[0].merge(tr[6])
        tr[0].paragraphs[0].add_run("Всего / Total").bold = True
        for cell, v in zip(tr[7:], [ctx["total_amount"], ctx["total_vat"], ctx["total_with_vat"]]):
            cell.paragraphs[0].add_run(v).bold = True

        d.add_paragraph(
            f"\nОбщая стоимость работ по акту: {ctx['total_with_vat']} ({ctx['words_ru']}), "
            f"в том числе НДС {ctx['total_vat']} ({ctx['vat_words_ru']})."
        )
        d.add_paragraph(
            f"Total cost of work according to the act: {ctx['total_with_vat']} ({ctx['words_en']}), "
            f"including VAT {ctx['total_vat']} ({ctx['vat_words_en']})."
        )
        if ctx["wells_str"]:
            d.add_paragraph(
                f"По результатам работ: на скважинах №№ {ctx['wells_str']} продолжить работы "
                "по оптимизации/обслуживанию. Стороны претензий не имеют."
            )

        d.add_paragraph("\nПОДПИСИ СТОРОН / SIGNATURES OF THE PARTIES")
        d.add_paragraph("Исполнитель / The Contractor: ____________ Яцкив А. П., Директор / Director")
        d.add_paragraph("Заказчик / The Customer: ____________ Israilov U. T., Председатель Правления")
        d.add_paragraph("____________ Cho Eun Sang, Первый Заместитель Председателя Правления")

        base_name = document.doc_number.replace("/", "-")
        (self.output_dir / "docx").mkdir(exist_ok=True)
        out = self.output_dir / "docx" / f"{base_name}.docx"
        d.save(str(out))
        return f"generated/docx/{out.name}"

    @staticmethod
    def _find_soffice() -> str:
        import shutil
        for c in ("soffice", "libreoffice"):
            p = shutil.which(c)
            if p:
                return p
        mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.exists(mac):
            return mac
        raise RuntimeError("LibreOffice (soffice) не найден — не могу сделать PDF из .docx")

    def generate_pdf_from_docx(self, document: Document) -> str:
        """PDF финансового акта = конвертация того же .docx через LibreOffice
        (гарантирует идентичность .docx и PDF — инвариант PARITY)."""
        docx_rel = self.generate_docx(document)
        docx_path = (Path("backend/static") / docx_rel).resolve()
        pdf_dir = self.output_dir / "pdf"
        pdf_dir.mkdir(exist_ok=True)
        # изолированный профиль: конвертация не падает, даже если LibreOffice уже открыт
        # (или запущены параллельные headless-конвертации)
        profile = (self.output_dir / "temp" / f"lo_{document.doc_number.replace('/', '-')}").resolve()
        result = subprocess.run(
            [self._find_soffice(), "-env:UserInstallation=file://" + str(profile),
             "--headless", "--convert-to", "pdf",
             "--outdir", str(pdf_dir.resolve()), str(docx_path)],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120,
        )
        pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.exists():
            err = result.stderr.decode("utf-8", "ignore")
            raise RuntimeError(f"LibreOffice не создал PDF: {err}")
        document.pdf_filename = f"generated/pdf/{pdf_path.name}"
        return document.pdf_filename

    def generate_invoice_pdf(self, document: Document) -> str:
        """PDF счёта-фактуры через xelatex/longtable.

        longtable штатно повторяет на продолжении только строку номеров колонок
        (\\endhead) и не «протекает» вертикальными линиями — то, что недостижимо
        в docx→LibreOffice. Контекст — общий с актом (_prepare_financial_act_context).
        Выгрузка .docx (generate_docx) остаётся отдельно, как редактируемый вариант.
        """
        from backend.services.daily_report_service import _get_latex_env, _compile_latex
        ctx = self._prepare_financial_act_context(document)
        # Формат номера СФ: «СФ-01/2026» (порядковый с нулём / год). Порядковый —
        # ведущие цифры из invoice_no («1-с» → 1), год — из даты акта.
        import re
        seq = re.match(r"\d+", str(ctx.get("invoice_no") or ""))
        seq_n = int(seq.group()) if seq else (ctx.get("act_no") or 1)
        year = (ctx.get("act_date") or "")[-4:]
        ctx["invoice_label"] = f"СФ-{seq_n:02d}/{year}"
        env = _get_latex_env()
        tex = env.get_template("financial_invoice.tex").render(**ctx)
        pdf = _compile_latex(tex, document.doc_number.replace("/", "-"))
        document.pdf_filename = f"generated/pdf/{pdf.name}"
        return document.pdf_filename

    def _prepare_reagent_expense_context(self, document: Document) -> dict:
        """Подготовить контекст для акта расхода реагентов"""
        # Форматируем даты
        period_start_str = document.period_start.strftime("%d.%m.%Y") if document.period_start else ""
        period_end_str = document.period_end.strftime("%d.%m.%Y") if document.period_end else ""

        # Номер акта (число) и месяц
        doc_num_parts = document.doc_number.split('-')
        act_num = doc_num_parts[-1] if len(doc_num_parts) > 0 else "1"
        act_month = document.metadata.get("act_month_name_ru", "")

        # Номер скважины
        well_number = document.metadata.get("well_number",
                                            str(document.well.number) if document.well and document.well.number else "")

        # Дата акта (последний день периода)
        act_date = period_end_str

        # Подсчёты
        total_injections = len(document.items)
        summary = document.metadata.get("summary_by_type", {})

        return {
            "theactnum": act_num,
            "theactmonth": act_month,
            "theactwell": well_number,
            "theactdate": act_date,
            "period_start_str": period_start_str,
            "period_end_str": period_end_str,
            "total_injections": total_injections,
            "summary_foam": summary.get("foam", 0),
            "summary_inhibitor": summary.get("inhibitor", 0),
            "company_executor": document.metadata.get("company_executor", "ООО «UNITOOL»"),
            "company_client": document.metadata.get("company_client", "СП ООО «Uz-Kor Gas Chemical»"),
            "field_name": document.metadata.get("field_name", "Сургил"),
        }

    def _render_latex_template(self, template_name: str, context: dict) -> str:
        """
        Рендерить LaTeX шаблон

        Args:
            template_name: Имя файла шаблона
            context: Контекст для Jinja2

        Returns:
            str: Рендеренный LaTeX код
        """
        template = self.latex_env.get_template(template_name)
        return template.render(**context)

    def _compile_latex_to_pdf(self, document: Document, latex_source: str) -> str:
        """
        Скомпилировать LaTeX в PDF

        Args:
            document: Объект Document
            latex_source: LaTeX код

        Returns:
            str: Полный путь к PDF файлу
        """
        # Имя файла без расширения
        base_name = f"{document.doc_number.replace('/', '-')}"

        # Временная директория
        temp_dir = self.output_dir / "temp"
        tex_file = temp_dir / f"{base_name}.tex"

        # Сохраняем .tex файл
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_source)

        # Компилируем в PDF (xelatex для поддержки русского)
        # Запускаем дважды для правильных ссылок
        for _ in range(2):
            result = subprocess.run(
                [_find_xelatex(), '-interaction=nonstopmode', f'-output-directory={temp_dir}', str(tex_file)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                cwd=str(temp_dir)
            )

            if result.returncode != 0:
                error_msg = result.stderr.decode('utf-8', errors='ignore')
                raise RuntimeError(f"Ошибка компиляции LaTeX: {error_msg}")

        # Перемещаем PDF в output
        temp_pdf = temp_dir / f"{base_name}.pdf"
        final_pdf = self.output_dir / "pdf" / f"{base_name}.pdf"

        if temp_pdf.exists():
            temp_pdf.rename(final_pdf)
        else:
            raise FileNotFoundError(f"PDF файл не создан: {temp_pdf}")

        # Очищаем временные файлы
        for ext in ['.aux', '.log', '.out', '.tex']:
            temp_file = temp_dir / f"{base_name}{ext}"
            if temp_file.exists():
                temp_file.unlink()

        return str(final_pdf)

    # ==========================================
    # ГЕНЕРАЦИЯ EXCEL
    # ==========================================

    def generate_excel(self, document: Document) -> str:
        """
        Сгенерировать Excel для документа

        Args:
            document: Объект Document

        Returns:
            str: Путь к сгенерированному Excel (относительно static/)
        """
        try:
            from openpyxl import Workbook, load_workbook
            from openpyxl.styles import Font, Alignment, Border, Side
        except ImportError:
            raise ImportError("Установите openpyxl: pip install openpyxl")

        # Получаем имя шаблона
        template_name = document.doc_type.excel_template_name

        if template_name and (self.excel_dir / template_name).exists():
            # Используем шаблон
            wb = load_workbook(self.excel_dir / template_name)
            ws = wb.active
        else:
            # Создаём с нуля
            wb = Workbook()
            ws = wb.active
            ws.title = "Акт"

        # Для акта расхода реагентов
        if document.doc_type.code == 'reagent_expense':
            self._fill_reagent_expense_excel(ws, document)

        # Сохраняем
        base_name = f"{document.doc_number.replace('/', '-')}"
        excel_path = self.output_dir / "excel" / f"{base_name}.xlsx"
        wb.save(excel_path)

        # Сохраняем путь в документ
        document.excel_filename = f"generated/excel/{excel_path.name}"

        return document.excel_filename

    def _fill_reagent_expense_excel(self, ws, document: Document):
        """Заполнить Excel для акта расхода реагентов"""
        from openpyxl.styles import Font, Alignment

        # Заголовок
        ws['A1'] = f"Акт №{document.doc_number}"
        ws['A1'].font = Font(bold=True, size=14)

        ws['A2'] = f"дозирования реагентов на скважине №{document.metadata.get('well_number', '')}"
        ws['A3'] = f"месторождения {document.metadata.get('field_name', 'Сургил')}"
        ws[
            'A4'] = f"в период с {document.period_start.strftime('%d.%m.%Y')} по {document.period_end.strftime('%d.%m.%Y')}"

        # Заголовки таблицы
        row = 6
        headers = ['№', 'Вид работ', 'Дата и время вброса', 'К-во', 'Тип реагента', 'Этап']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')

        # Данные
        row += 1
        for item in document.items:
            ws.cell(row=row, column=1, value=item.line_number)
            ws.cell(row=row, column=2, value=item.work_type)
            ws.cell(row=row, column=3, value=item.event_time_str)
            ws.cell(row=row, column=4, value=item.quantity)
            ws.cell(row=row, column=5, value=item.reagent_name)
            ws.cell(row=row, column=6, value=item.stage)
            row += 1

        # Итого
        row += 1
        ws.cell(row=row, column=1, value=f"Всего вбросов — {len(document.items)}")
        ws.cell(row=row, column=1).font = Font(bold=True)

        # Ширина колонок
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 35
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 8
        ws.column_dimensions['E'].width = 20
        ws.column_dimensions['F'].width = 25