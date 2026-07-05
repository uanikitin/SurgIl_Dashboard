"""Однократно: превращает реальный .docx пользователя в docxtpl-шаблон.

Вход:  /Users/volodymyrnikitin/Downloads/Акт финансовый  январь_3 (3).docx
Выход: backend/documents/templates/docx/financial_act_template.docx

Вставляет jinja-теги docxtpl:
- шапка: № {{act_no}} от {{act_date}}, период {{period_from}}..{{period_to}}
- таблица работ: строка-цикл {%tr for r in rows %} ... {%tr endfor %}
- итоги: {{total_amount}}/{{total_vat}}/{{total_with_vat}}
- решение+пропись: {{wells_str}}, {{words_ru}}/{{words_en}}, {{vat_words_ru}}/{{vat_words_en}}
"""
from pathlib import Path
from docx import Document as Dx

SRC = "/Users/volodymyrnikitin/Downloads/Акт финансовый  январь_3 (3).docx"
OUT = Path("backend/documents/templates/docx/financial_act_template.docx")


def set_para(paragraph, text: str):
    """Заменить текст абзаца одним run (формат первого run сохраняется)."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def find_para(cell, needle):
    for p in cell.paragraphs:
        if needle in p.text:
            return p
    return None


def del_row(table, idx):
    tr = table.rows[idx]._tr
    tr.getparent().remove(tr)


def main():
    d = Dx(SRC)

    # ── TABLE 0: шапка + преамбула (RU c0 / EN c1) ──
    t0 = d.tables[0]
    ru, en = t0.rows[0].cells[0], t0.rows[0].cells[1]
    p = find_para(ru, "№ 3");  set_para(p, "№ {{ act_no }} от {{ act_date }}") if p else None
    p = find_para(en, "# 3");  set_para(p, "# {{ act_no }} dated {{ act_date }}") if p else None

    pre_ru, pre_en = t0.rows[1].cells[0], t0.rows[1].cells[1]
    p = find_para(pre_ru, "в период с")
    if p:
        set_para(p,
            "Мы, нижеподписавшиеся, представители заказчика СП ООО «Uz-Kor Gas Chemical» "
            "{{ header_customer }}, действующие на основании Устава, с одной стороны, и "
            "представитель подрядчика ООО «UNITOOL» {{ header_contractor }}, действующий на "
            "основании Устава, с другой стороны, составили настоящий Акт о том, что "
            "специалистами Исполнителя в период с {{ period_from }} по {{ period_to }} "
            "были выполнены следующие работы:")
    p = find_para(pre_en, "during the period")
    if p:
        set_para(p,
            "We, undersigned, representatives of Customer JV «Uz-Kor Gas Chemical» LLC "
            "{{ header_customer_en }}, on the one part, and representative of Contractor "
            "«UNITOOL» LLC {{ header_contractor_en }}, on the other part, have hereby executed "
            "this Act that specialists of Contractor have performed the following works during "
            "the period {{ period_from }} to {{ period_to }}.")

    # ── TABLE 1: строки работ → цикл docxtpl ──
    # ВАЖНО: тег {%tr ... %} удаляет всю строку, где он стоит. Поэтому for и endfor —
    # в ОТДЕЛЬНЫХ строках-обёртках вокруг повторяемой строки данных.
    import copy
    t1 = d.tables[1]
    tpl_idx = 3            # строка-данные (Адаптация/85) — станет повторяемой
    last_idx = len(t1.rows) - 1  # «Всего/Total»
    keep = {0, 1, tpl_idx, last_idx}

    # итоговая строка
    total_cells = t1.rows[last_idx].cells
    for ci, tag in ((8, "{{ total_amount }}"), (9, "{{ total_vat }}"), (10, "{{ total_with_vat }}")):
        set_para(total_cells[ci].paragraphs[0], tag)

    # строка-данные: плоские теги (без tr)
    c = t1.rows[tpl_idx].cells
    tags = {
        0: "{{ r.n }}", 2: "{{ r.name }}", 3: "{{ r.well }}", 4: "{{ r.period }}",
        5: "{{ r.unit }}", 6: "{{ r.qty }}", 7: "{{ r.price }}", 8: "{{ r.amount }}",
        9: "{{ r.vat }}", 10: "{{ r.total }}",
    }
    for ci, tag in tags.items():
        set_para(c[ci].paragraphs[0], tag)

    # удалить прочие строки данных
    for idx in sorted((i for i in range(len(t1.rows)) if i not in keep), reverse=True):
        del_row(t1, idx)

    # вставить строки-обёртки for/endfor вокруг строки-данных (теперь она по индексу 2)
    data_row = t1.rows[2]
    for_tr = copy.deepcopy(data_row._tr)
    end_tr = copy.deepcopy(data_row._tr)
    data_row._tr.addprevious(for_tr)
    data_row._tr.addnext(end_tr)
    # теперь строки: 0 header,1 num,2 FOR,3 data,4 ENDFOR,5 total
    def clear_and_tag(row, tag):
        for ci, cell in enumerate(row.cells):
            set_para(cell.paragraphs[0], "")
        set_para(row.cells[0].paragraphs[0], tag)
    clear_and_tag(t1.rows[2], "{%tr for r in rows %}")
    clear_and_tag(t1.rows[4], "{%tr endfor %}")

    # ── TABLE 2: решение + пропись ──
    t2 = d.tables[2]

    def rebuild_cell(cell, lines):
        """Очистить ячейку и заполнить заново (убирает дубли-абзацы)."""
        for par in list(cell.paragraphs):
            par._p.getparent().remove(par._p)
        for txt in lines:
            cell.add_paragraph(txt)

    # решение — обе ячейки (RU и EN) с нуля, чтобы не осталось статичных дублей
    rebuild_cell(t2.rows[0].cells[0], [
        "По результатам работ было принято следующее решение:",
        "{{ decision_continue }} {{ decision_stop }}",
        "Выполненные объёмы работ полностью удовлетворяют условиям Контракта. "
        "Стороны претензий и замечаний друг к другу не имеют.",
    ])
    rebuild_cell(t2.rows[0].cells[1], [
        "Based on the results of the work, the following decision was made:",
        "{{ decision_continue_en }} {{ decision_stop_en }}",
        "The scope of work performed fully satisfies the terms of the Contract. "
        "The parties have no claims and comments to each other.",
    ])
    sum_ru = find_para(t2.rows[1].cells[0], "Общая стоимость")
    if sum_ru:
        set_para(sum_ru,
            "Общая стоимость работ по акту: {{ total_with_vat }} ({{ words_ru }}), "
            "в том числе НДС {{ total_vat }} ({{ vat_words_ru }}).")
    sum_en = find_para(t2.rows[1].cells[1], "Total cost")
    if sum_en:
        set_para(sum_en,
            "Total cost of work according to the act: {{ total_with_vat }} ({{ words_en }}), "
            "including VAT {{ total_vat }} ({{ vat_words_en }}).")

    # ── TABLE 3: блок подписей → циклы по сторонам (несколько подписантов) ──
    t3 = d.tables[3]

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    Wsig = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def rebuild_sig_cell(cell, header, loop_var):
        for par in list(cell.paragraphs):           # очистить ячейку
            par._p.getparent().remove(par._p)
        lines = [
            header,
            "", "",                                  # место под заголовком → черта ниже
            "{%p for s in " + loop_var + " %}",
            "________________________________",
            "{{ s.name_ru }} ({{ s.name_en }})",
            "{{ s.position_ru }} / {{ s.position_en }}",
            "", "",                                  # интервал между подписантами
            "{%p endfor %}",
        ]
        for txt in lines:
            p = cell.add_paragraph(txt)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER   # подпись по центру

    rebuild_sig_cell(t3.rows[1].cells[0], "The Contractor (Исполнитель)", "sign_contractor")
    rebuild_sig_cell(t3.rows[1].cells[1], "The Customer (Заказчик)", "sign_customer")

    # убрать ВСЕ границы таблицы подписей (вертикальные линии под таблицей работ)
    tblPr3 = t3._tbl.tblPr
    for old in tblPr3.findall(Wsig + "tblBorders"):
        tblPr3.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge); e.set(qn("w:val"), "nil"); borders.append(e)
    tblPr3.append(borders)
    for row in t3.rows:
        for tc in row._tr.findall(Wsig + "tc"):
            tcPr = tc.find(Wsig + "tcPr")
            if tcPr is not None:
                for tb in tcPr.findall(Wsig + "tcBorders"):
                    tcPr.remove(tb)

    # ── Единый шрифт 10pt в таблице работ ──
    from docx.shared import Pt, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for row in t1.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    # ── Фиксированные ширины колонок (см), сумма ≤ usable 26.7 (landscape A4) ──
    # период может быть межмесячным («09.03-21.04.2026», 16 симв.) → колонка 4 пошире
    widths_cm = [0.8, 0.7, 5.5, 1.6, 3.1, 2.2, 1.3, 2.6, 2.9, 2.7, 2.9]
    tbl = t1._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))  # 1 см = 567 twips
        grid.append(gc)
    for row in t1.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                cell.width = Cm(widths_cm[ci])

    # ── Совпадение ширин шапки и тела: в шапке «№» объединить на колонки 0+1,
    #    «Наименование» сделать одинарной (как в строках данных). Иначе граница
    #    между колонками 1 и 2 в шапке и теле стоит в разных местах. ──
    Wns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def set_span(tc, val):
        tcPr = tc.find(Wns + "tcPr")
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
        for old in tcPr.findall(Wns + "gridSpan"):
            tcPr.remove(old)
        gs = OxmlElement("w:gridSpan"); gs.set(qn("w:val"), str(val))
        tcPr.insert(0, gs)

    for ri in (0, 1):  # строка заголовков + строка нумерации «1 2 3 …»
        tcs = t1.rows[ri]._tr.findall(Wns + "tc")
        if len(tcs) >= 2:
            set_span(tcs[0], 2)  # № → колонки 0+1
            set_span(tcs[1], 1)  # Наименование → одинарная (колонка 2)
    # повтор заголовка (названия + номера колонок) на каждой новой странице при разрыве
    for ri in (0, 1):
        tr = t1.rows[ri]._tr
        trPr = tr.find(Wns + "trPr")
        if trPr is None:
            trPr = OxmlElement("w:trPr"); tr.insert(0, trPr)
        for old in trPr.findall(Wns + "tblHeader"):
            trPr.remove(old)
        trPr.append(OxmlElement("w:tblHeader"))

    # запрет разрыва строк (чтобы merged-ячейки «Наименование» не «протекали»
    # вертикальными линиями при разрыве таблицы между страницами)
    for row in t1.rows:
        tr = row._tr
        trPr = tr.find(Wns + "trPr")
        if trPr is None:
            trPr = OxmlElement("w:trPr"); tr.insert(0, trPr)
        if trPr.find(Wns + "cantSplit") is None:
            trPr.append(OxmlElement("w:cantSplit"))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(OUT))
    print("Saved template:", OUT)


if __name__ == "__main__":
    main()
