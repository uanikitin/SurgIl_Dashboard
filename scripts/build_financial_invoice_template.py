"""Строит docxtpl-шаблон Счёта-фактуры ИЗ ШАБЛОНА АКТА.

Акт рендерится корректно (таблица работ полной ширины), поэтому СФ делаем как
копию шаблона акта с заменой шапки на реквизиты счёта-фактуры и удалением блока
решения. Таблица работ (T1) и подписи (T3) переиспользуются как есть.

Вход:  financial_act_template.docx (собран build_financial_act_template.py)
Выход: backend/documents/templates/docx/financial_invoice_template.docx
"""
from pathlib import Path
from docx import Document as Dx

ACT_TPL = "backend/documents/templates/docx/financial_act_template.docx"
OUT = Path("backend/documents/templates/docx/financial_invoice_template.docx")

SERVICE = ("на оказание услуг по оптимизации эксплуатации газовых скважин с применением "
           "твёрдых пенообразующих реагентов на скважинах месторождения Сургил")

BANK_CONTRACTOR = [
    "ИСПОЛНИТЕЛЬ / CONTRACTOR:",
    "«UNITOOL» LLC",
    "TIN 309222928",
    "Kungrad district, Republic of Karakalpakstan,",
    "Kungrad region, Azatliq MFY, G'a'rezsizlik street, house 11",
    "Bank details:",
    "В/а: 20208000905483697001",
    "in OPERU JSCB «Kapitalbank»",
    "Bank code: 00974",
    "TIN: 203591761  OKPO: 17763371  OKED: 64190",
]
BANK_CUSTOMER = [
    "ЗАКАЗЧИК / CLIENT:",
    "JV «Uz-Kor Gas Chemical» LLC",
    "Republic of Uzbekistan, Republic of Karakalpakstan, Nukus city, Karatau,",
    "G'arezsizlik MFY, To'rtko'l guzari street, building 121",
    "Address of Tashkent office: 100128, Zulfiyakhonim str. 112",
    "Tel/fax: +998-78-129-29-00. E-mail: info@uz-kor.com",
    "Bank details:",
    "B/a: 20214000904704378001 in CJSC «KDB Bank Uzbekistan» Tashkent",
    "Bank code: 00842  OKONH: 11232  TIN: 300 829 145",
]


def rebuild_cell(cell, lines):
    for par in list(cell.paragraphs):
        par._p.getparent().remove(par._p)
    for txt in lines:
        cell.add_paragraph(txt)


def main():
    d = Dx(ACT_TPL)
    t0 = d.tables[0]  # шапка/преамбула акта → шапка СФ

    # r0: заголовок + номер + контракт + описание услуги (обе колонки)
    rebuild_cell(t0.rows[0].cells[0], [
        "СЧЁТ-ФАКТУРА / INVOICE",
        "№ {{ invoice_no }} от {{ act_date }}",
        "согласно Контракту № {{ contract_ref }} / to Contract № {{ contract_ref }}",
        SERVICE,
    ])
    rebuild_cell(t0.rows[0].cells[1], [""])
    # r1: банковские реквизиты (Исполнитель | Заказчик)
    rebuild_cell(t0.rows[1].cells[0], BANK_CONTRACTOR)
    rebuild_cell(t0.rows[1].cells[1], BANK_CUSTOMER)

    # удалить таблицу решения (T2) — в счёте-фактуре её нет
    t2 = d.tables[2]._tbl
    t2.getparent().remove(t2)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(OUT))
    print("Saved invoice template:", OUT)


if __name__ == "__main__":
    main()
