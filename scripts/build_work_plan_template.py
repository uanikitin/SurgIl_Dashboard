"""Однократно: авторская сборка docxtpl-шаблона «План работ» (сценарий ОПТИМИЗАЦИЯ)
по эталону PDF «План работ скважина 74.pdf».

Выход: backend/documents/templates/docx/work_plan_template.docx

Динамика (docxtpl-теги):
- шапка СОГЛАСОВАНО / УТВЕРЖДАЮ: {{ sog_position }}/{{ sog_org }}/{{ sog_name }}/{{ sog_date }}
  и {{ utv_* }} (всё редактируемо);
- в блок УТВЕРЖДАЮ вставляются картинки {{ signature }} и {{ seal }} (печать 4×4 см,
  подпись — с сохранением пропорций); подставляются генератором как InlineImage;
- номер скважины {{ well_no }} (в заголовке и строке 1 Таблицы 1.1);
- значения Таблицы 1.1 — {{ t.* }}.

Статика (Цель, Таблица 2.1, Таблица 3.1, раздел 4) вшита ниже, идентична эталону.

Запуск:  .venv/bin/python scripts/build_work_plan_template.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("backend/documents/templates/docx/work_plan_template.docx")

FONT = "Times New Roman"
UZK = 'JV "Uz-Kor Gas Chemical" LLC'
UNI = 'ООО "UNITOOL"'
BOTH = f"{UZK}\n{UNI}"


# ─────────────────────────── низкоуровневые помощники ───────────────────────────

def _set_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)


def para(container, text="", *, size=11, bold=False, italic=False,
         align=None, space_after=4, space_before=0, keep_next=False):
    """Добавить абзац с единым форматированием. Многострочный текст (\\n) → line-break.
    keep_next=True — не отрывать от следующего элемента (для заголовков таблиц)."""
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if keep_next:
        pf.keep_with_next = True
    lines = text.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        _set_run(r, size=size, bold=bold, italic=italic)
        if i < len(lines) - 1:
            r.add_break()
    return p


def cell_text(cell, text, *, size=10, bold=False, align=None):
    """Записать текст в ячейку (первый абзац), поддержка \\n."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        _set_run(r, size=size, bold=bold)
        if i < len(lines) - 1:
            r.add_break()


def _borders(table, val):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), val)
        if val != "none":
            e.set(qn("w:sz"), "4"); e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)


def _set_col_widths_pct(table, widths_pct):
    """Ширины колонок в процентах от ширины таблицы. Сумма должна быть 100."""
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Таблица 100% ширины страницы
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")  # 5000 = 100%
    tblPr.insert(0, tblW)
    # Фиксированная раскладка
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # Ширины колонок в процентах (50 pct units = 1%)
    grid = tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for pct in widths_pct:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(pct * 50)))  # 50 units = 1%
        grid.append(gc)
    # Ширины ячеек
    for row in table.rows:
        for ci, c in enumerate(row.cells):
            if ci < len(widths_pct):
                tcPr = c._tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:type"), "pct")
                tcW.set(qn("w:w"), str(int(widths_pct[ci] * 50)))
                tcPr.append(tcW)


def _repeat_header_rows(table, n=1):
    """Установить повтор первых n строк таблицы при переносе на новую страницу."""
    for i in range(n):
        tr = table.rows[i]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        hdr = OxmlElement("w:tblHeader")
        trPr.append(hdr)


# ──────────────────────────────── шапка ────────────────────────────────

def build_header(doc):
    """Две колонки: СОГЛАСОВАНО (Uz-Kor) | УТВЕРЖДАЮ (UNITOOL, с печатью и подписью)."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders(t, "none")
    _set_col_widths_pct(t, [50, 50])  # 50% + 50% = 100%
    left, right = t.rows[0].cells

    def fill(cell, title, pos_tag, org_tag, name_tag, date_tag, pad_lines=0):
        cell.text = ""
        # Заголовок — по центру
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(p.add_run(title), size=11, bold=True)
        # Остальные строки — по левому краю
        para(cell, pos_tag, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=0)
        para(cell, org_tag, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
        for _ in range(pad_lines):
            para(cell, "", size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
        # Увеличенные отступы между строками ФИО и подписи
        para(cell, f"_______________  {name_tag}", size=10,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6)
        para(cell, f"____  _____________  {date_tag}", size=10,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)

    fill(left, "СОГЛАСОВАНО", "{{ sog_position }}", "{{ sog_org }}",
         "{{ sog_name }}", "{{ sog_date }}")
    # Печать floating — не нужны пустые строки
    fill(right, "УТВЕРЖДАЮ", "{{ utv_position }}", "{{ utv_org }}",
         "{{ utv_name }}", "{{ utv_date }}", pad_lines=0)


# ──────────────────────────────── Таблица 1.1 ────────────────────────────────

TABLE11_ROWS = [
    ("№ скважины / Well No.", "{{ t.well_no }}"),
    ("Диаметр эксплуатационной колонны, мм / Production casing diameter, mm", "{{ t.prod_casing_diam }}"),
    ("Глубина спуска эксплуатационной колонны, м / Depth of running production casing, m", "{{ t.prod_casing_depth }}"),
    ("Текущий забой, м / Current bottomhole, m", "{{ t.bottomhole }}"),
    ("Горизонт / Horizon", "{{ t.horizon }}"),
    ("Интервалы перфорации, м / Perforation intervals, m", "{{ t.perforation }}"),
    ("Диаметр НКТ, мм / Diameter of tubing, mm", "{{ t.tubing_diam }}"),
    ("Глубина башмака НКТ, м / Tubing shoe depth, m", "{{ t.tubing_shoe }}"),
    ("Глубина пакера, м / Packer depth, m", "{{ t.packer }}"),
    ("Глубина переводника, м / Adapter depth, m", "{{ t.adapter }}"),
    ("Глубина непрохода шаблона, м / Depth of pattern stuck, m", "{{ t.pattern_stuck }}"),
    ("Диаметр штуцера, мм / Choke diameter, mm", "{{ t.choke }}"),
    ("Трубное давление, кгс/см² / Tubing pressure, kgf/cm²", "{{ t.p_tube }}"),
    ("Затрубное давление, кгс/см² / Annulus pressure, kgf/cm²", "{{ t.p_annulus }}"),
    ("Шлейфное давление, кгс/см² / Flowline pressure, kgf/cm²", "{{ t.p_flowline }}"),
    ("Дебит газа, тыс.м³/сут / Gas rate, ths.m³/day", "{{ t.q_gas }}"),
    ("Статическое давление, кгс/см² / Static pressure, kgf/cm²", "{{ t.p_static }}"),
    ("Пластовое давление, кгс/см² / Reservoir pressure, kgf/cm²", "{{ t.p_reservoir }}"),
    ("Периодичность продувки, час / Blowing frequency, hour", "{{ t.purge_freq }}"),
    ("Длительность продувки, мин / Blowing duration, min", "{{ t.purge_dur }}"),
    ("Результаты последних ГДИ/ГКИ/ОПП (прилагаются к заявке) / "
     "Results of latest GDS/GCS/PLT (attached to the application)", "{{ t.gdi }}"),
]


def build_table11(doc):
    para(doc, "Таблица 1.1: Параметры скважины / Well data", size=11, italic=True,
         space_before=6, space_after=2, keep_next=True)
    # +1 заголовок + 1 нумерация + данные
    t = doc.add_table(rows=2 + len(TABLE11_ROWS), cols=2)
    _borders(t, "single")
    _set_col_widths_pct(t, [70, 30])  # 70% показатель, 30% значение
    # Строка 0: заголовки
    cell_text(t.rows[0].cells[0], "Показатель / Mark", size=10, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[0].cells[1], "Значение / Value", size=10, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    # Строка 1: нумерация колонок (повторяется при переносе)
    cell_text(t.rows[1].cells[0], "1", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[1], "2", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _repeat_header_rows(t, n=2)  # заголовок + нумерация
    # Данные
    for i, (label, tag) in enumerate(TABLE11_ROWS, start=2):
        cell_text(t.rows[i].cells[0], label, size=10)
        cell_text(t.rows[i].cells[1], tag, size=10)


# ──────────────────────────────── Таблица 2.1 ────────────────────────────────

TABLE21 = [
    ('Шлюз устьевой ШУ-50-35-550 ТУ У 29.5–40846045–002: 2018 / '
     'Wellhead launcher SHU-50-35-550', "компл/kit", "1", UNI),
    ('Пенообразующий реагент / Foaming reagents', "шт/piece", "в ас-те", UNI),
    ('Тестовый сепаратор ФГС – 50–10,0/RS / Mobile test separator (MTS)', "компл/kit", "1", UNI),
    ('Автономные СИ с функцией логирования / Stand-alone measuring instruments '
     'with logging function', "компл/kit", "2", UNI),
    ('Система мониторинга регистрации параметров работы скважины / System for '
     'monitoring and recording well operating parameters SMOD', "компл/kit", "1", UNI),
    ('Емкость для слива отсепарированного скважинного флюида объемом не менее 1 м3 '
     'с возможностью подключения дренажной линии к БРС 2" / Tank for drainage of '
     'separated well fluid with the volume not less than 1 m3 with the possibility '
     'of connecting the drainage line to the 2" quick-release connection.', "шт/piece", "1", UNI),
]


def build_table21(doc):
    para(doc, "Таблица 2.1: Перечень оборудования и материалов / Equipment and materials list",
         size=11, italic=True, space_before=6, space_after=2, keep_next=True)
    # +1 заголовок + 1 нумерация + данные
    t = doc.add_table(rows=2 + len(TABLE21), cols=5)
    _borders(t, "single")
    # № 5%, описание 55%, ед 9%, к-во 11%, ответственный 20%
    _set_col_widths_pct(t, [5, 55, 9, 11, 20])
    # Строка 0: заголовки
    hdr = ["№", "Перечень материалов и оборудования / Materials and equipment list",
           "ед", "к-во / quantity", "Ответственный / Responsibility"]
    for ci, h in enumerate(hdr):
        cell_text(t.rows[0].cells[ci], h, size=10, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    # Строка 1: нумерация колонок
    for ci in range(5):
        cell_text(t.rows[1].cells[ci], str(ci + 1), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _repeat_header_rows(t, n=2)
    # Данные
    for i, (name, unit, qty, resp) in enumerate(TABLE21, start=2):
        cell_text(t.rows[i].cells[0], f"{i-1}.", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(t.rows[i].cells[1], name, size=10)
        cell_text(t.rows[i].cells[2], unit, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(t.rows[i].cells[3], qty, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(t.rows[i].cells[4], resp, size=10)


# ──────────────────────────────── Таблица 3.1 ────────────────────────────────
# ("stage", текст) — строка-подзаголовок; (номер, "RU / EN", ответственный) — работа.

TABLE31 = [
    ("stage", "Подготовительные работы. Этап 1 / Preparatory work. Stage 1"),
    ("1", "Проведение ЗАКАЗЧИКОМ инструктажа по охране труда и пожарной безопасности согласно "
          "нормативным документам Республики Узбекистан. / Conducting a briefing on labor and fire "
          "safety protection in accordance with regulatory documents of the Republic of Uzbekistan.", UZK),
    ("2", "Ознакомление представителей ИСПОЛНИТЕЛЯ с материалами скважины, режимом и параметрами "
          "работы скважины. / Familiarization of representatives of the Contractor with Well materials, "
          "the mode and parameters of the Well.", UZK),
    ("3", "Ознакомление представителей ИСПОЛНИТЕЛЯ с устройством системы сбора и подготовки скважинной "
          "продукции. / Familiarization of representatives of the contractor with the device of the "
          "system of collecting and preparing well products.", UZK),
    ("4", "Подготовка скважины к проведению технологических операций: установка рабочей площадки для "
          "обслуживания скважин согласно требованиям нормативных документов Республики Узбекистан; "
          "проверка работоспособности запорной арматуры и герметичности фонтанной арматуры; проверка "
          "исправности контрольно-измерительных приборов (манометров) на буфере, затрубном пространстве, "
          "за штуцером и на выходе из промыслового коллектора. / Well preparation for technological "
          "operations: installation of the Well service site in accordance with the requirements of "
          "regulatory documents of the Republic of Uzbekistan; checking the performance of shut-off "
          "valves and tightness of fountain reinforcement; checking the serviceability of instrumentation "
          "and measuring instruments (pressure gauges) on the buffer, casing space, for Choke and at the "
          "exit from the commercial manifold.", UZK),
    ("5", "Установить на скважину измерительное оборудование / Install measuring equipment on the well", UNI),
    ("6", "Осуществить сбор технологических данных перед проведением полевых работ (температура, давление "
          "в НКТ, затрубном пространстве, давление в линии за штуцером) не менее 3 циклов. / Collect "
          "process data prior to field work (temperature, tubing and annulus pressure, line pressure "
          "behind the choke) for at least 3 cycles.", UNI),
    ("7", "На основании полученных данных произвести расчет текущего дебита и анализ режимов работы "
          "скважины / Based on the obtained data, calculate the current flow rate and analyze the well "
          "operation modes", UNI),
    ("8", "Произвести оценочный расчет периодичности и дозировки химреагентов с учетом фактических режимов "
          "работы скважины, выполнить оценку и прогнозирование ожидаемых результатов от применения "
          "технологии оптимизации работы скважины с использованием ТПАВ. / To make an estimation "
          "calculation of frequency and dosage of chemical reagents taking into account actual well "
          "operation modes, to make an estimation and forecasting of expected results from application of "
          "well operation optimization technology with the use of surfactants.", UNI),
    ("9", "Согласовать скважину для дальнейшего проведения работ. / Coordinate the well for further work.", BOTH),
    ("stage", "Проведение полевых работ по адаптации технологии к условиям скважины. Этап 2 / "
              "Conduct field work to adapt the technology to the well conditions. Stage 2"),
    ("10", "Установка Шлюзового устройства ШУ-50-35-550 на буфере скважины и опрессовка его рабочим "
           "давлением. / Installation Wellhead Launcher SHU-50-35-550 at the well buffer and testing its "
           "working pressure.", UNI),
    ("11", "Обеспечить наличие на скважине средств мониторинга и контроля рабочих параметров скважины на "
           "НКТ и затрубном пространстве. / Ensure availability of monitoring and control of well "
           "operating parameters on tubing and annular space at the well.", UNI),
    ("12", "На основании полученных данных со средств мониторинга, произвести предварительный подбор "
           "химических реагентов для обеспечения вспенивания и выноса пластового флюида. При необходимости "
           "провести отбор проб пластового флюида с помощью тестового сепаратора. / Based on the data "
           "received from the monitoring equipment, make a preliminary selection of chemical reagents to "
           "ensure foaming and removal of reservoir fluid. If necessary, conduct sampling of reservoir "
           "fluid using a test separator.", UNI),
    ("13", "Доставка на скважину химреагентов / Delivery to the well of chemical reagents", UNI),
    ("14", "Введение расчетной дозы пенообразователя в скважину. / Injection of the calculated dose of the "
           "foaming agent in the well.", UNI),
    ("15", "Производить наблюдение контрольных параметров работы скважины (дебит газа, давления на буфере, "
           "в затрубном пространстве, после штуцера и на выходе из промыслового коллектора). / Monitor "
           "well control parameters (gas flow rate, buffer pressure, pressure in the annulus, after the "
           "choke and at the outlet from the field manifold).", BOTH),
    ("16", "При падении трубного давления, установившегося после отработки скважины согласно п. 15, до "
           "давления близкого к давлению выкидной линии, произвести равномерное введение шашек "
           "пенообразователя в скважину с периодичностью и количеством рассчитанными на основании анализа "
           "данных с контрольно-измерительных приборов. / When the pipe pressure, established after "
           "working out the well according to p. 15, drops to a pressure close to the pressure of the "
           "discharge line, to make a uniform introduction of the foamers into the well with the frequency "
           "and quantity calculated on the basis of analyzing the data from the control-measuring devices.", UNI),
    ("17", "Снятие контрольных параметров (давления на буфере, в затрубном пространстве, после штуцера и на "
           "выходе из промыслового коллектора) не реже трех раз в течение суток. При необходимости провести "
           "замер текущего дебита газа и состав флюида. / Measurement of control parameters (pressure at "
           "the buffer, in the annulus, after the choke and at the outlet from the field manifold) at least "
           "three times per day. If necessary, measure the current gas flow rate and fluid composition.", BOTH),
    ("18", "По результатам п. 17 проводятся изменения в периодичности и количестве вводимого "
           "пенообразователя. / According to the results of p. 17 changes in the frequency and quantity of "
           "the introduced blowing agent are carried out.", UNI),
    ("19", "Произвести выбор оптимального химреагента, учитывая изменения устьевых параметров скважины "
           "после вброса реагента. / Select the optimal chemical reagent, taking into account changes in "
           "wellhead parameters after reagent injection.", UNI),
    ("20", "После изменения периодичности и количества вводимого пенообразователя проводятся работы "
           "согласно п.п. 15–18 до обеспечения оптимальных условий работы скважины. / After changing the "
           "periodicity and quantity of the injected blowing agent the works according to p.p. 15-18 are "
           "carried out until optimal conditions of well operation are achieved.", UNI),
    ("21", "Предоставление ежедневного отчета ЗАКАЗЧИКУ о выполненных работах, с содержанием полной "
           "информации о применяемой технике и времени ее работы, количестве и времени применения "
           "используемого реагента, сведения о выполненных замерах и их результаты, сведения о персонале, "
           "задействованном при выполнении работ. / Submission of a daily report to the CUSTOMER on the "
           "work performed, containing full information on the equipment used and its operating time, the "
           "amount and time of application of the reagent used, information on the measurements made and "
           "their results, information on the personnel involved in the work.", UNI),
    ("23", "Предоставить рекомендации по применению дополнительных реагентов для оптимизации работы "
           "скважины (ингибиторы коррозии, поглотители сероводорода, растворители солеотложений, "
           "деэмульгаторы, депрессаторы и пр.) / Provide recommendations on the use of additional reagents "
           "to optimize well performance (corrosion inhibitors, hydrogen sulfide absorbers, salt deposit "
           "solvents, demulsifiers, depressants, etc.).", ""),
    ("24", "Сбор замечаний ЗАКАЗЧИКА к технологии с целью их устранения при дальнейшем ее использовании. / "
           "Collection of the CUSTOMER's comments on the technology in order to eliminate them for future "
           "use.", UNI),
    ("25", "Провести анализ и оценку экономической целесообразности постоянного дозирования реагента на "
           "скважине с учетом выбранных режимов дозирования. / Analyze and evaluate the economic "
           "feasibility of constant dosing of the reagent at the well, taking into account the selected "
           "modes of dosing.", ""),
    ("26", "Согласовать скважину для дальнейшего проведения работ. / Approve continued work on the well.", BOTH),
    ("stage", "Проведение полевых работ для обеспечения стабильной работы скважины. Этап 3 / "
              "Conducting field work to ensure well stable operation. Stage 3"),
    ("27", "Обеспечить наличие на скважине шлюзового устройства для дозирования реагента, средств "
           "мониторинга и контроля рабочих параметров скважины на НКТ, линии и затрубном пространстве / "
           "Ensure that there is a Wellhead launcher for reagent dosing, means of monitoring and control "
           "of well operating parameters on tubing, line and downhole space", UNI),
    ("28", "Обеспечить непрерывный мониторинг и анализ текущих рабочих параметров скважины, с ежеминутной "
           "фиксацией и логированием данных, обеспечить возможность синхронизации данных со всех типов "
           "измерительных устройств / Provide continuous monitoring and analysis of the current operating "
           "parameters of the well, with minute-by-minute recording and logging of data, provide the "
           "ability to synchronize data from all types of measuring devices", UNI),
    ("29", "Производить вброс пенообразователя, согласно выбранного режима. Контролировать изменения "
           "устьевых параметров. Не допускать снижение буферного давления до уровня давления в линии. При "
           "необходимости изменить периодичность дозирования. / To inject the blowing agent according to "
           "the selected mode. Monitor changes in wellhead parameters. Do not allow the buffer pressure to "
           "drop to the line pressure level. If necessary, change the dosing frequency.", UNI),
    ("30", "При необходимости провести замер рабочих параметров скважины с помощью мобильного сепаратора с "
           "целью определения качества пены, текущих фактических дебитов газа и флюида. Замер произвести на "
           "рабочем штуцере скважины в атмосферных условиях, при этом в сепараторе создавая давление "
           "сепарации равное шлейфному давлению. / If necessary, measure well operating parameters using a "
           "mobile separator to determine the quality of foam, current actual gas and fluid flow rates. "
           "Measurement should be performed at the working well connection under atmospheric conditions, "
           "while creating a separation pressure equal to the plume pressure in the separator.", BOTH),
    ("31", "Обеспечить предоставление ежедневной сводки ЗАКАЗЧИКУ о выполненных работах, с содержанием "
           "полной информации о применяемой технике и времени ее работы, количестве и времени применения "
           "используемого реагента, сведения о выполненных замерах и их результаты, сведения о персонале, "
           "задействованном при выполнении работ. / Provide the CUSTOMER with a daily summary of the work "
           "performed, containing full information on the equipment used and its operating time, the amount "
           "and time of application of the reagent used, information on the measurements made and their "
           "results, information on the personnel involved in the execution of the work.", UNI),
    ("32", "Обеспечить отображение типа и количества использованных реагентов в суточных сводках. Вести "
           "учет использованных реагентов / Ensure that the type and quantity of reagents used are "
           "displayed in daily summaries. Maintain records of reagents used", UNI),
    ("33", "В конце отчетного периода предоставить Акт использованных реагентов / At the end of the "
           "reporting period, provide the Statement of used reagents", BOTH),
    ("34", "В конце отчетного периода согласовать продолжение работ на следующий отчетный период / At the "
           "end of the reporting period, agree on the continuation of work for the next reporting period", BOTH),
]


def build_table31(doc):
    para(doc, "Таблица 3.1: Перечень работ, выполняемых на скважине / List of works performed on the well",
         size=11, italic=True, space_before=6, space_after=2, keep_next=True)
    # +1 заголовок + 1 нумерация + данные
    t = doc.add_table(rows=2 + len(TABLE31), cols=3)
    _borders(t, "single")
    # № 5%, описание 75%, ответственный 20%
    _set_col_widths_pct(t, [5, 75, 20])
    # Строка 0: заголовки
    hdr = ["№", "Вид работ / Type of work", "Ответственный / Responsibility"]
    for ci, h in enumerate(hdr):
        cell_text(t.rows[0].cells[ci], h, size=10, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    # Строка 1: нумерация колонок
    for ci in range(3):
        cell_text(t.rows[1].cells[ci], str(ci + 1), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _repeat_header_rows(t, n=2)
    # Данные
    for i, item in enumerate(TABLE31, start=2):
        row = t.rows[i]
        if item[0] == "stage":
            merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            cell_text(merged, item[1], size=10, bold=True)
        else:
            num, text, resp = item
            cell_text(row.cells[0], num, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_text(row.cells[1], text, size=10)
            cell_text(row.cells[2], resp, size=10)


# ──────────────────────────────── Раздел 4 ────────────────────────────────

SAFETY = [
    ("h", "4 МЕРОПРИЯТИЯ ПО БЕЗОПАСНОМУ ВЕДЕНИЮ РАБОТ / SAFETY MEASURES"),
    ("p", "При разработке месторождений, на которых намечается реализация технологии физико-химического "
          "воздействия, для обеспечения безопасности населения необходимо руководствоваться действующими "
          "законами, постановлениями и положениями, Правилами безопасности в нефтяной и газовой "
          "промышленности. / During the development of fields, where it is planned to implement the "
          "technology of physical-chemical impact to ensure the safety of the population must be guided by "
          "applicable laws, decrees and regulations, Safety Rules in the oil and gas industry."),
    ("sh", "4.1 Меры безопасности при работе / Safety precautions at work"),
    ("p", "4.1.1 В процессе проведения работ соблюдать требования действующего законодательства, иных "
          "нормативно-правовых актов и нормативных документов Республики Узбекистан в области промышленной "
          "безопасности и охраны труда. / In the course of work, comply with the requirements of the "
          "current legislation, other regulatory legal acts and regulatory documents of the Republic of "
          "Uzbekistan in the field of industrial safety and labor protection."),
    ("p", "4.1.2 Персонал, привлекаемый для выполнения данных работ, должен быть обучен на курсах по "
          "вопросам промышленной безопасности и охране труда и при выполнении работ на месторождении "
          "Сургил должен при себе иметь удостоверение единого образца, установленного уполномоченным "
          "органом в области промышленной безопасности. / Personnel engaged for this work shall be trained "
          "at the courses on industrial safety and labor protection issues and when performing works at "
          "the Surgil field shall have a certificate of a single sample established by the authorized body "
          "in the field of industrial safety."),
    ("p", "4.1.3 Все применяемые реагенты не токсичны, пожаро- и взрывобезопасны. / All reagents used are "
          "non-toxic, fire and explosion safe."),
    ("p", "4.1.4 Категорически запрещается использовать реагенты для нужд, не связанных с их прямой целью. / "
          "It is strictly forbidden to use reagents for purposes other than their direct purpose."),
    ("p", "4.1.5 При появлении признаков отравления (головной боли, головокружения, тошноты, рвоты, потери "
          "аппетита), следует обратиться к врачу и сообщить руководителю работ. / In case of signs of "
          "poisoning (headache, dizziness, nausea, vomiting, loss of appetite), consult a doctor and inform "
          "the work manager."),
    ("sh", "4.2 Охрана окружающей среды / Environmental protection"),
    ("p", "4.2.1 После выполнения работ необходимо / After the work is done, it is necessary to:"),
    ("li", "Принимать все меры безопасности, необходимые для защиты окружающей среды, оберегая окружающий "
           "воздух, поверхностные и подземные воды, почвы и грунты, недра, животный и растительный мир от "
           "неблагоприятных воздействий, вызванных действиями, и сводя к минимуму ущерб, который может "
           "повлечь за собой подобные действия. / Take all safety measures necessary to protect the "
           "environment by protecting the surrounding air, surface and underground water, soils and "
           "subsoil, fauna and flora from adverse effects caused by the activities and minimizing the "
           "damage that may result from such activities."),
    ("li", "Обеспечить своевременный вывоз и утилизацию коммунальных и производственных отходов. / Ensure "
           "that municipal and industrial waste is removed and disposed of in a timely manner."),
    ("li", "Обеспечить очистку устьевого оборудования, шахты, рекультивацию площадки, планировку и сдачу "
           "Заказчику по акту. / Ensure cleaning of wellhead equipment, shaft, site reclamation, layout "
           "and handover to the Customer under the act."),
]


def build_safety(doc):
    for kind, text in SAFETY:
        if kind == "h":
            para(doc, text, size=12, bold=True, space_before=10, space_after=4)
        elif kind == "sh":
            para(doc, text, size=11, bold=True, space_before=6, space_after=3)
        elif kind == "li":
            p = para(doc, "•  " + text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)
            p.paragraph_format.left_indent = Cm(0.8)
        else:
            para(doc, text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)


# ──────────────────────────────── сборка ────────────────────────────────

def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2.0); s.right_margin = Cm(1.5)

    build_header(doc)

    para(doc, "ПЛАН РАБОТ / PLAN OF WORK", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    para(doc, "по оптимизации работы скважины {{ well_no }} с применением твердых пенообразующих реагентов / "
              "on optimization of well {{ well_no }} operation with the use of solid foaming reagents",
         size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    para(doc, "Цель работ: подбор рецептуры, концентрации и режимов дозирования пенообразователя для "
              "обеспечения стабильной работы газодобывающей скважины путем удаления жидкости из ствола "
              "скважины. / Purpose of work: selection of formulation, concentration and dosing regimes of "
              "blowing agent to ensure stable operation of gas production well by removing fluid from the "
              "wellbore.", size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    para(doc, "1 ГЕОЛОГО - ТЕХНИЧЕСКИЕ ДАННЫЕ / GEOLOGICAL AND ENGINEERING DATA",
         size=12, bold=True, space_before=6, space_after=4)
    build_table11(doc)

    para(doc, "2 МАТЕРИАЛЫ И ОБОРУДОВАНИЕ / MATERIALS AND EQUIPMENT",
         size=12, bold=True, space_before=10, space_after=4)
    build_table21(doc)

    para(doc, "3 ПОРЯДОК ПРОВЕДЕНИЯ РАБОТ / OPERATING PROCEDURE",
         size=12, bold=True, space_before=10, space_after=4)
    build_table31(doc)

    build_safety(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("Saved template:", OUT)


if __name__ == "__main__":
    main()
